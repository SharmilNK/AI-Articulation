"""
Generates a formatted Word document (.docx) from the synthesized digest.
Sections mirror the 5 use cases: conversation starters, recent updates,
core concepts, roadblocks, mock conversations, and dev standards.
"""

import os
import logging
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "output")


def _ensure_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def _set_heading_style(para, level: int = 1):
    """Apply heading style without relying on template styles."""
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x53, 0xA1)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def _add_section_rule(doc: Document):
    """Add a thin horizontal rule after a heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0xA1)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    _add_section_rule(doc)


def _add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _add_label(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    val_run = p.add_run(value)
    val_run.font.size = Pt(10)


def _add_bullet(doc: Document, text: str, indent_level: int = 1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * indent_level)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10)


def _add_dialogue_line(doc: Document, speaker: str, line: str, is_consultant: bool):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    speaker_run = p.add_run(f"{speaker}: ")
    speaker_run.bold = True
    speaker_run.font.size = Pt(10)
    if is_consultant:
        speaker_run.font.color.rgb = RGBColor(0x1A, 0x53, 0xA1)
    else:
        speaker_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    line_run = p.add_run(line)
    line_run.font.size = Pt(10)
    if is_consultant:
        line_run.italic = True


def build_document(digest: dict, today_str: str) -> str:
    _ensure_output_dir()
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AI Articulation Daily Digest")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x1A, 0x53, 0xA1)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(today_str)
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = tagline.add_run(
        "Your daily brief to sound sharp in any AI conversation — executive boardroom or data science deep-dive."
    )
    t_run.italic = True
    t_run.font.size = Pt(11)
    doc.add_paragraph()

    # ── Section 1: Conversation Starters ──────────────────────────────────────
    _add_h1(doc, "1. Today's AI Conversation Starters")
    intro = doc.add_paragraph(
        "Use these to confidently open or redirect any discussion about AI."
    )
    intro.runs[0].font.size = Pt(10)
    intro.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("conversation_starters", []):
        _add_h2(doc, item.get("topic", ""))
        _add_label(doc, "Opening hook", item.get("hook", ""))
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("Talking points:").bold = True
        p.runs[0].font.size = Pt(10)
        for pt in item.get("talking_points", []):
            _add_bullet(doc, pt)
        _add_label(doc, "Source", item.get("source", ""))
        doc.add_paragraph()

    # ── Section 2: Recent Updates & Opinions ──────────────────────────────────
    _add_h1(doc, "2. Recent AI Updates & Your Opinion")
    intro2 = doc.add_paragraph(
        "Stay current. Know what changed. Have a take ready."
    )
    intro2.runs[0].font.size = Pt(10)
    intro2.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("recent_updates_opinions", []):
        _add_h2(doc, item.get("headline", ""))
        _add_label(doc, "In plain English", item.get("plain_english", ""))
        _add_label(doc, "Your take", item.get("opinion", ""))
        _add_label(doc, "Why it matters", item.get("why_it_matters", ""))
        doc.add_paragraph()

    # ── Section 3: Core AI/ML Concepts ────────────────────────────────────────
    _add_h1(doc, "3. Core AI/ML Concepts — Explained Simply")
    intro3 = doc.add_paragraph(
        "Master these and you'll command any room — boardroom or lab."
    )
    intro3.runs[0].font.size = Pt(10)
    intro3.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("core_concepts", []):
        _add_h2(doc, item.get("concept", ""))
        _add_label(doc, "Definition", item.get("one_liner", ""))
        _add_label(doc, "Think of it like", item.get("analogy", ""))
        _add_label(doc, "Why it's hot right now", item.get("why_relevant_now", ""))
        doc.add_paragraph()

    # ── Section 4: AI Roadblocks & Solutions ──────────────────────────────────
    _add_h1(doc, "4. AI Roadblocks & What the Industry Is Doing")
    intro4 = doc.add_paragraph(
        "Know the obstacles — and how smart teams are getting around them."
    )
    intro4.runs[0].font.size = Pt(10)
    intro4.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("roadblocks_solutions", []):
        _add_h2(doc, item.get("problem", ""))
        _add_label(doc, "Context", item.get("context", ""))
        _add_label(doc, "Industry response", item.get("industry_response", ""))
        _add_label(doc, "How to frame it", item.get("consultant_angle", ""))
        doc.add_paragraph()

    # ── Section 5: Mock Conversations ─────────────────────────────────────────
    _add_h1(doc, "5. Mock Conversations — Practice These")

    # 5a Executive
    exec_conv = digest.get("mock_exec_conversation", {})
    _add_h2(doc, f"With a Business Executive — {exec_conv.get('scenario', '')}")
    for turn in exec_conv.get("dialogue", []):
        speaker = turn.get("speaker", "")
        line = turn.get("line", "")
        is_consultant = speaker.lower() == "consultant"
        _add_dialogue_line(doc, speaker, line, is_consultant)

    kp = exec_conv.get("key_phrases_used", [])
    if kp:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("Key phrases to remember: ").bold = True
        p.runs[0].font.size = Pt(10)
        p.add_run(", ".join(f'"{ph}"' for ph in kp)).font.size = Pt(10)
    doc.add_paragraph()

    # 5b Data Scientist
    ds_conv = digest.get("mock_ds_conversation", {})
    _add_h2(doc, f"With a Senior Data Scientist — {ds_conv.get('scenario', '')}")
    for turn in ds_conv.get("dialogue", []):
        speaker = turn.get("speaker", "")
        line = turn.get("line", "")
        is_consultant = speaker.lower() == "consultant"
        _add_dialogue_line(doc, speaker, line, is_consultant)

    kp_ds = ds_conv.get("key_phrases_used", [])
    if kp_ds:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("Key phrases to remember: ").bold = True
        p.runs[0].font.size = Pt(10)
        p.add_run(", ".join(f'"{ph}"' for ph in kp_ds)).font.size = Pt(10)
    doc.add_paragraph()

    # ── Section 6: Dev & GitHub Standards ─────────────────────────────────────
    _add_h1(doc, "6. GitHub & Dev Standards — What's New")
    intro6 = doc.add_paragraph(
        "Know the tools your engineering teams use — it builds instant credibility."
    )
    intro6.runs[0].font.size = Pt(10)
    intro6.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    std = digest.get("dev_standards_update", {})
    _add_h2(doc, std.get("standard", "Developer Standards Update"))
    _add_label(doc, "What it is", std.get("what_it_is", ""))
    _add_label(doc, "Why AI startups care", std.get("why_ai_startups_care", ""))
    _add_label(doc, "Smart thing to say", std.get("talking_point", ""))
    doc.add_paragraph()

    # ── Section 7: AI Governance Frameworks ───────────────────────────────────
    _add_h1(doc, "7. AI Governance Frameworks — Permanent Reference")
    intro7 = doc.add_paragraph(
        "The industry-standard frameworks your clients are already using or being asked about. "
        "Revisit these daily until they're second nature."
    )
    intro7.runs[0].font.size = Pt(10)
    intro7.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("governance_frameworks", []):
        _add_h2(doc, f"{item.get('framework', '')}  —  {item.get('owner', '')}")
        _add_label(doc, "What it is", item.get("one_liner", ""))
        pillars = item.get("core_pillars", [])
        if pillars:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("Core pillars: ").bold = True
            p.runs[0].font.size = Pt(10)
            p.add_run(" · ".join(pillars)).font.size = Pt(10)
        _add_label(doc, "How companies use it", item.get("how_companies_use_it", ""))
        _add_label(doc, "Say this in a meeting", item.get("consultant_talking_point", ""))
        doc.add_paragraph()

    # ── Section 8: AI Strategy Frameworks ─────────────────────────────────────
    _add_h1(doc, "8. AI Strategy Frameworks — Permanent Reference")
    intro8 = doc.add_paragraph(
        "Strategic frameworks from the leading voices in enterprise AI. "
        "Use these when advising on AI roadmaps, build-vs-buy decisions, and organisational readiness."
    )
    intro8.runs[0].font.size = Pt(10)
    intro8.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for item in digest.get("ai_strategy_frameworks", []):
        _add_h2(doc, f"{item.get('framework', '')}  [{item.get('source', '')}]")
        _add_label(doc, "Core idea", item.get("core_idea", ""))
        _add_label(doc, "Real-world application", item.get("real_world_application", ""))
        checklist = item.get("deployment_checklist", [])
        if checklist:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("Deployment checklist:").bold = True
            p.runs[0].font.size = Pt(10)
            for step in checklist:
                _add_bullet(doc, step)
        _add_label(doc, "Say this in a C-suite meeting", item.get("consultant_talking_point", ""))
        doc.add_paragraph()

    # ── Section 9: Enterprise AI Adoption Toolkit ─────────────────────────────
    _add_h1(doc, "9. Enterprise AI Adoption Toolkit — Permanent Reference")
    intro9 = doc.add_paragraph(
        "The frameworks and models your clients need when deploying AI at scale. "
        "Use these to lead maturity assessments, build business cases, and manage organisational change."
    )
    intro9.runs[0].font.size = Pt(10)
    intro9.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    TOPIC_ICONS = {
        "AI Maturity Assessment": "Maturity Assessment",
        "Roadmap Definition": "Roadmap Definition",
        "Executive Advisory": "Executive Advisory",
        "Business Case Development": "Business Case Development",
        "Change Management": "Change Management",
        "Enterprise AI Adoption": "Enterprise AI Adoption",
    }

    for item in digest.get("enterprise_ai_toolkit", []):
        topic = item.get("topic", "")
        source = item.get("source", "")
        _add_h2(doc, f"{topic}  [{source}]")
        _add_label(doc, "Framework / Model", item.get("framework_or_model", ""))
        _add_label(doc, "What it covers", item.get("description", ""))
        steps = item.get("key_steps_or_levels", [])
        if steps:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            p.add_run("Key steps / levels:").bold = True
            p.runs[0].font.size = Pt(10)
            for step in steps:
                _add_bullet(doc, step)
        _add_label(doc, "How you use this with a client", item.get("how_consultant_uses_it", ""))
        _add_label(doc, "Open with this", item.get("client_talking_point", ""))
        doc.add_paragraph()

    # Footer note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run(
        "Generated by AI-Articulation | Sources: Lenny's Newsletter, Spill the GPTea, The Unwind AI, "
        "Intuitive Autonomy, Enterprise AI Executive, Bridge by Adi Agrawal, Tomorrow Toolbox, Maven, "
        "Conventional Commits | Governance: Microsoft, Google, NIST | Strategy: a16z, MIT Sloan, Stanford HAI"
    )
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    f_run.italic = True

    filename = f"AI_Articulation_Digest_{date.today().strftime('%Y-%m-%d')}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    logger.info("Document saved: %s", filepath)
    return filepath
